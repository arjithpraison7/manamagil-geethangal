import docx
import csv

# Load the docx file
index_doc = docx.Document('Manamakizh geetham Index.docx')

rows = []
# Extract tables if present
for table in index_doc.tables:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

# Fallback: If no tables, try paragraphs as before
if not rows:
    for para in index_doc.paragraphs:
        text = para.text.strip()
        if text:
            # Split by tab or multiple spaces, or just keep as one column if no delimiter
            if '\t' in text:
                row = text.split('\t')
            elif '  ' in text:
                row = [col.strip() for col in text.split('  ') if col.strip()]
            else:
                row = [text]
            rows.append(row)

# Write to CSV
with open('manamakizh_index.csv', 'w', encoding='utf-8', newline='') as f:
    writer = csv.writer(f)
    writer.writerows(rows)

print('Index extracted to manamakizh_index.csv')
